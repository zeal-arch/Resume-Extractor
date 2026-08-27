"""
src/parsers/docx_parser.py
──────────────────────────
Responsible for extracting content from DOCX files:
  - Visible text from paragraphs AND table cells (via python-docx).
  - Embedded hyperlink URIs from the document's relationship store.

Why table cells matter
──────────────────────
Many resume templates use Word tables for layout (contact-info row, skills grid,
two-column experience blocks). doc.paragraphs silently skips table cell content;
we must iterate doc.tables separately to capture it.

Why relationships matter
────────────────────────
DOCX hyperlinks (mailto:, https://) are stored in the part's .rels XML,
not in the visible text. python-docx exposes these via doc.part.rels.
"""

from docx import Document


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract all visible text from a DOCX file.

    Sources
    -------
    1. doc.paragraphs  — standard body text, headings, list items.
    2. doc.tables      — table cells joined per row with ' | ' separators,
                         so the text can still be parsed as a line.

    Returns an empty string on failure (error is printed to stdout).
    """
    try:
        doc = Document(docx_path)
        parts: list[str] = []

        # ── 1. Paragraph text ──────────────────────────────────────────────
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # ── 2. Table cell text ─────────────────────────────────────────────
        # Cells in the same row are joined with ' | ' so they appear on one
        # line and don't confuse section-header detection.
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append("  |  ".join(row_cells))

        return "\n".join(parts)

    except Exception as exc:
        print(f"[docx_parser] Error reading '{docx_path}': {exc}")
        return ""


def extract_embedded_links_from_docx(docx_path: str) -> list[str]:
    """
    Extract hyperlink target URIs from the DOCX relationship store.

    DOCX files store hyperlinks in an XML relationships file (.rels).
    python-docx exposes these as doc.part.rels, keyed by relationship ID.
    We collect any target_ref that starts with 'mailto:' or 'http'.

    Returns a list of URI strings.
    Returns an empty list on failure.
    """
    uris: list[str] = []
    try:
        doc = Document(docx_path)
        for rel in doc.part.rels.values():
            target = rel.target_ref
            if target and (target.startswith('mailto:') or target.startswith('http')):
                uris.append(target)
    except Exception as exc:
        print(f"[docx_parser] Warning: could not read hyperlinks from '{docx_path}': {exc}")

    return uris
